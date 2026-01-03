import os
import re
from copy import deepcopy
from pathlib import Path
import numpy as np
import pandas as pd
from sklearn.metrics import mean_squared_error
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# =========================
# ===== USER INPUTS =======
# =========================

TEMPLATE_DOC_PATH = r"C:/Users/satya/Downloads/Template.docx"
OUTPUT_DIR        = r"C:/Work/WeatherCast/Weekly Advisory/Washim"

FORECAST_START = pd.Timestamp("2025-12-29")
FORECAST_END   = pd.Timestamp("2026-01-04")

locations = [
    {"name": "Netansa", "obs_file": r"C:/Users/satya/Downloads/WC2507004_Netansa.csv", "model_file": r"C:/Users/satya/Downloads/open-meteo-20.25N76.75E535m.csv"},
    {"name": "Sakra", "obs_file": r"C:/Users/satya/Downloads/WC2507001_Sakra.csv", "model_file": r"C:/Users/satya/Downloads/open-meteo-20.25N77.25E487m.csv"},
    {"name": "Malegaon", "obs_file": r"C:/Users/satya/Downloads/WC2507002_Malegaon.csv", "model_file": r"C:/Users/satya/Downloads/open-meteo-20.25N77.00E515m.csv"},    
]

bulletin_numbers = {
    "Netansa": 13,
    "Sakra": 13,
    "Malegaon": 13,
}

manual_rainfall = {k: "--" for k in bulletin_numbers}
manual_wind     = {k: "--" for k in bulletin_numbers}

GENERATE_MODE = "per_location"

# =========================
# ====== HELPERS ==========
# =========================

def add_date_suffix(day: int) -> str:
    if 10 <= day % 100 <= 20:
        suf = "th"
    else:
        suf = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
    return f"{day}{suf}"

def issued_on_str(dt: pd.Timestamp) -> str:
    return f"{add_date_suffix(dt.day)} {dt.strftime('%B')}"

def valid_till_str(dt: pd.Timestamp) -> str:
    return f"{add_date_suffix(dt.day)} {dt.strftime('%B')} {dt.year}"

def set_cell_text_preserve_format(cell, new_text):
    if cell.paragraphs and cell.paragraphs[0].runs:
        original_format = deepcopy(cell.paragraphs[0].runs[0]._element.rPr)
        para = cell.paragraphs[0]
        for r in para.runs:
            r.text = ""
        run = para.add_run(str(new_text))
        if original_format is not None:
            run._element.insert(0, original_format)
    else:
        cell.text = str(new_text)

def _replace_span_in_runs(para: Paragraph, start: int, end: int, replacement: str) -> bool:
    if not para.runs:
        return False
    pos = 0
    s_idx = s_off = e_idx = e_off = None
    for idx, run in enumerate(para.runs):
        nxt = pos + len(run.text)
        if s_idx is None and start <= nxt:
            s_idx = idx
            s_off = start - pos
        if s_idx is not None and end <= nxt:
            e_idx = idx
            e_off = end - pos
            break
        pos = nxt
    if s_idx is None or e_idx is None:
        return False
    before = para.runs[s_idx].text[:s_off]
    after  = para.runs[e_idx].text[e_off:]
    if s_idx == e_idx:
        para.runs[s_idx].text = before + replacement + after
    else:
        para.runs[s_idx].text = before + replacement
        for j in range(s_idx + 1, e_idx):
            para.runs[j].text = ""
        para.runs[e_idx].text = after
    return True

RE_BULLETIN   = re.compile(r"Bulletin\s*:?\s*(\d+)")
RE_ISSUED_ON  = re.compile(r"(Issued on\s+)(\d{1,2}(?:st|nd|rd|th)\s+[A-Za-z]+)")
RE_VALID_TILL = re.compile(r"(Valid till.*?of\s+)(\d{1,2}(?:st|nd|rd|th)\s+[A-Za-z]+\s+\d{4})")

def replace_in_paragraph(para: Paragraph, pattern: re.Pattern, repl_text: str, group_index: int) -> bool:
    combined = "".join(run.text for run in para.runs)
    m = pattern.search(combined)
    if not m:
        return False
    start, end = m.span(group_index)
    return _replace_span_in_runs(para, start, end, repl_text)

# =========================
# === TEMPLATE MAPPING ====
# =========================

def build_location_blocks(doc: Document):
    """Map each 'Weather Advisory' block to location name + paragraph span."""
    paras = doc.paragraphs
    n = len(paras)
    title_idxs = [i for i, p in enumerate(paras) if "Weather Advisory" in p.text]
    blocks, order = {}, []

    for ti, start_idx in enumerate(title_idxs):
        end_idx = title_idxs[ti + 1] if ti + 1 < len(title_idxs) else n
        loc_name = None
        for j in range(start_idx, end_idx):
            t = paras[j].text.strip()
            if t.startswith("Location:"):
                loc_name = t.split("Location:", 1)[1].strip()
                break
        if loc_name:
            blocks[loc_name] = (start_idx, end_idx)
            order.append(loc_name)

    return blocks, order

def tables_in_range(doc: Document, start_para: Paragraph, end_para: Paragraph):
    """Return tables whose XML position lies between the two paragraph XML elements."""
    def _el_idx(elem):
        return elem.getparent().index(elem)
    s = _el_idx(start_para._element)
    e = _el_idx(end_para._element)
    out = []
    for tbl in doc.tables:
        t = _el_idx(tbl._element)
        if s <= t < e:
            out.append(tbl)
    return out
# =========================
# ======= FORECAST ========
# =========================

def compute_forecast_for_location(obs_file, model_file, start_date, end_date) -> pd.DataFrame:
    obs = pd.read_csv(obs_file)
    obs["date_time"] = pd.to_datetime(obs["date_time"], dayfirst=True, errors="coerce")
    obs["Date"] = obs["date_time"].dt.date

    daily_obs = obs.groupby("Date").agg({
        "temperature": ["max", "min"],
        "humidity": ["max", "min"],
        "precipitation": "sum"
    }).reset_index()
    daily_obs.columns = ["Date", "TMAX", "TMIN", "RHMAX", "RHMIN", "RAINFALL"]
    daily_obs["Date"] = pd.to_datetime(daily_obs["Date"])

    model = pd.read_csv(model_file, skiprows=3)
    model["time"] = pd.to_datetime(model["time"], errors="coerce")
    model["Date"] = pd.to_datetime(model["time"].dt.date)

    cols = model.columns.tolist()
    temp_cols = [c for c in cols if "temperature_2m_" in c.lower()]
    rh_cols   = [c for c in cols if "relative_humidity_2m_" in c.lower()]
    rain_cols = [c for c in cols if "rain_" in c.lower() and "(mm)" in c.lower()]

    model_cols = {"TMAX": temp_cols, "TMIN": temp_cols, "RHMAX": rh_cols, "RHMIN": rh_cols, "RAINFALL": rain_cols}
    agg_map = {"TMAX": "max", "TMIN": "min", "RHMAX": "max", "RHMIN": "min", "RAINFALL": "sum"}

    best_models, bias = {}, {}
    for param, cands in model_cols.items():
        rmse_dict, bias_dict = {}, {}
        for c in cands:
            try:
                dm = model.groupby("Date")[c].agg(agg_map[param]).reset_index()
                dm.columns = ["Date", "ModelVal"]
                merged = pd.merge(daily_obs[["Date", param]], dm, on="Date", how="inner").dropna()
                if not merged.empty:
                    rmse_dict[c] = np.sqrt(mean_squared_error(merged[param], merged["ModelVal"]))
                    bias_dict[c] = (merged["ModelVal"] - merged[param]).mean()
            except Exception:
                pass
        if rmse_dict:
            best = min(rmse_dict, key=rmse_dict.get)
            best_models[param] = best
            bias[param] = bias_dict[best]

    ecmwf = [c for c in rain_cols if "ecmwf_ifs025" in c.lower()]
    if ecmwf:
        best_models["RAINFALL"] = ecmwf[0]
        bias["RAINFALL"] = 0.0

    dates = pd.date_range(start_date, end_date, freq="D")
    out = pd.DataFrame({"Date": dates})

    for param in ["TMAX", "TMIN", "RHMAX", "RHMIN", "RAINFALL"]:
        col = best_models.get(param)
        if not col:
            continue
        df = model.groupby("Date")[col].agg(agg_map[param]).reset_index()
        df.columns = ["Date", f"Raw_{param}"]
        df = df[df["Date"].isin(dates)]
        df[f"Corrected_{param}"] = df[f"Raw_{param}"] - bias.get(param, 0.0)
        out = pd.merge(out, df[["Date", f"Corrected_{param}"]], on="Date", how="left")

    return out

# =========================
# === DAILY MARATHI =======
# =========================

def _band_str(hi, delta):
    hi = round(float(hi), 1)
    lo = max(0.0, round(hi - delta, 1))
    lo_txt = "00" if abs(lo - 0.0) < 1e-9 else f"{lo:.1f}"
    return f"{lo_txt}-{hi:.1f}"

def rainfall_to_text(rainfall, date_str):
    if pd.isna(rainfall):
        return [date_str, "डेटा उपलब्ध नाही"]
    val = float(rainfall)
    if round(val, 1) == 0.0:
        text = "हवामान मुख्यत्वे कोरडे राहील."
    elif 0.01 <= val <= 5:
        text = f"कधीकधी ढगाळ, हलका पाऊस होऊ शकतो ({_band_str(val, 0.5)} मि.मी.)"
    elif 5 < val <= 15:
        text = f"मुख्यत्वे ढगाळ, हलका पाऊस होऊ शकतो ({_band_str(val, 1.0)} मि.मी.)"
    elif 15 < val <= 30:
        text = f"मुख्यत्वे ढगाळ, मध्यम पाऊस होऊ शकतो ({_band_str(val, 1.0)} मि.मी.)"
    elif 30 < val <= 45:
        text = f"मुख्यत्वे ढगाळ, जास्त पाऊस होऊ शकतो ({_band_str(val, 1.0)} मि.मी.)"
    elif 45 < val <= 60:
        text = f"मुख्यत्वे ढगाळ, जास्त पाऊस होऊ शकतो ({_band_str(val, 2.0)} मि.मी.)"
    elif 60 < val <= 125:
        text = f"मुख्यत्वे ढगाळ, खूप जास्त पाऊस होऊ शकतो ({_band_str(val, 5.0)} मि.मी.)"
    elif 125 < val <= 245:
        text = f"मुख्यत्वे ढगाळ, अतिशय जास्त पाऊस होऊ शकतो ({_band_str(val, 10.0)} मि.मी.)"
    else:
        text = f"मुख्यत्वे ढगाळ, प्रचंड पाऊस होऊ शकतो ({_band_str(val, 20.0)} मि.मी.)"
    return [date_str, text]

# =========================
# === WEEKLY LABELS =======
# =========================

_NBSP = u"\u00A0"

def _temp_range_text(series: pd.Series) -> str:
    if series.dropna().empty:
        return "माहिती उपलब्ध नाही"
    lo = float(series.min()); hi = float(series.max())
    return f"{lo:.1f}°C ते {hi:.1f}°C"

def _rh_range_text(series: pd.Series) -> str:
    if series.dropna().empty:
        return "माहिती उपलब्ध नाही"
    lo = int(round(float(series.min()))); hi = int(round(float(series.max())))
    return f"{lo}% ते {hi}%"

def _normalize_label(lbl: str) -> str:
    return re.sub(r"[\s\-\(\)%]+", "", lbl.strip().lower())

LABEL_MAP = {
    # English template labels
    "t-max": "tmax", "tmax": "tmax", "t-maxc": "tmax",
    "t-min": "tmin", "tmin": "tmin", "t-minc": "tmin",
    "rain": "rain", "rainfall": "rain", "rainfallmm": "rain",
    "windspeed&direction": "wind", "windspeeddirection": "wind", "wind": "wind",
    "rhmax": "rhmax", "maxrh": "rhmax", "rhmax%": "rhmax",
    "rhmin": "rhmin", "minrh": "rhmin", "rhmin%": "rhmin",

    # Marathi labels
    "कमालतापमान": "tmax",
    "किमानतापमान": "tmin",
    "पाऊस": "rain",
    "वाऱ्याचावेगआणिदिशा": "wind",
    "कमालसापेक्षआर्द्रता": "rhmax",
    "किमानसापेक्षआर्द्रता": "rhmin",
}

OUTPUT_LABELS_MARATHI = {
    "tmax": "कमाल तापमान",
    "tmin": "किमान तापमान",
    "rain": "पाऊस",
    "wind": "वाऱ्याचा वेग आणि दिशा",
    "rhmax": "कमाल सापेक्ष आर्द्रता",
    "rhmin": "किमान सापेक्ष आर्द्रता",
}

def _classify_label(raw: str):
    key = _normalize_label(raw)
    return LABEL_MAP.get(key)


# =========================
# ====== FILL BLOCK =======
# =========================

# (continued on next message if you want me to avoid cut-off)
def fill_location_block(doc: Document, block_range, loc_name, bulletin_no, forecast_df,
                        start_dt, end_dt, mr, mw):
    s_idx, e_idx = block_range
    paras = doc.paragraphs

    # 1) Header lines
    for i in range(s_idx, e_idx):
        p = paras[i]
        txt = p.text
        if "Bulletin" in txt:
            replace_in_paragraph(p, RE_BULLETIN, str(bulletin_no), group_index=1)
        elif "Valid for next 7 days, Issued on" in txt:
            replace_in_paragraph(p, RE_ISSUED_ON, issued_on_str(start_dt), group_index=2)
        elif txt.strip().startswith("Location:"):
            if p.runs:
                fmt = deepcopy(p.runs[0]._element.rPr)
                for r in p.runs:
                    r.text = ""
                rr = p.add_run(f"Location: {loc_name}")
                if fmt is not None:
                    rr._element.insert(0, fmt)
            else:
                p.text = f"Location: {loc_name}"

    # 2) Find the three tables inside this block
    start_para = paras[s_idx]
    end_para   = paras[e_idx - 1] if e_idx - 1 < len(paras) else paras[-1]
    tbls = tables_in_range(doc, start_para, end_para)

    top_tbl, daily_tbl, weekly_tbl = None, None, None

    for t in tbls:
        try:
            if t.cell(0, 0).text.strip() == "Block":
                top_tbl = t
                break
        except Exception:
            pass

    for t in tbls:
        if t is top_tbl:
            continue
        try:
            left_col = [row.cells[0].text.strip() for row in t.rows[1:]]
            if any(re.match(r"\d{2}/\d{2}/\d{4}$", s) for s in left_col):
                daily_tbl = t
                break
        except Exception:
            pass

    for t in tbls:
        if t not in (top_tbl, daily_tbl):
            weekly_tbl = t
            break

    # 3) Top table (numbers)
    if top_tbl is not None:
        for row in top_tbl.rows[:2]:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, RE_VALID_TILL, valid_till_str(end_dt), group_index=2)

        if not forecast_df.empty and len(top_tbl.rows) >= 7:
            dates_str = forecast_df["Date"].dt.strftime("%d/%m/%y").tolist()

            def _safe_round(x, nd=1, as_int=False):
                if pd.isna(x): return ""
                return int(round(x, 0)) if as_int else round(float(x), nd)

            for i in range(1, min(len(top_tbl.rows[1].cells), 1 + len(dates_str))):
                set_cell_text_preserve_format(top_tbl.rows[1].cells[i], dates_str[i - 1])
            for i in range(1, min(len(top_tbl.rows[2].cells), 1 + len(dates_str))):
                set_cell_text_preserve_format(
                    top_tbl.rows[2].cells[i],
                    _safe_round(forecast_df.loc[i - 1, "Corrected_RAINFALL"], 1)
                )
            for i in range(1, min(len(top_tbl.rows[3].cells), 1 + len(dates_str))):
                set_cell_text_preserve_format(
                    top_tbl.rows[3].cells[i],
                    _safe_round(forecast_df.loc[i - 1, "Corrected_TMAX"], 1)
                )
            for i in range(1, min(len(top_tbl.rows[4].cells), 1 + len(dates_str))):
                set_cell_text_preserve_format(
                    top_tbl.rows[4].cells[i],
                    _safe_round(forecast_df.loc[i - 1, "Corrected_TMIN"], 1)
                )
            for i in range(1, min(len(top_tbl.rows[5].cells), 1 + len(dates_str))):
                set_cell_text_preserve_format(
                    top_tbl.rows[5].cells[i],
                    _safe_round(forecast_df.loc[i - 1, "Corrected_RHMAX"], as_int=True)
                )
            for i in range(1, min(len(top_tbl.rows[6].cells), 1 + len(dates_str))):
                set_cell_text_preserve_format(
                    top_tbl.rows[6].cells[i],
                    _safe_round(forecast_df.loc[i - 1, "Corrected_RHMIN"], as_int=True)
                )

    # 4) Daily Marathi narrative
    if daily_tbl is not None and not forecast_df.empty:
        for i, row in enumerate(forecast_df.itertuples(index=False)):
            if i + 1 >= len(daily_tbl.rows): break
            date_str = row.Date.strftime("%d/%m/%Y")
            _, rain_desc = rainfall_to_text(getattr(row, 'Corrected_RAINFALL', np.nan), date_str)
            set_cell_text_preserve_format(daily_tbl.rows[i + 1].cells[0], date_str)
            set_cell_text_preserve_format(daily_tbl.rows[i + 1].cells[1], rain_desc)

    # 5) Weekly summary Marathi
    if weekly_tbl is not None and not forecast_df.empty:
        tmax_txt  = _temp_range_text(forecast_df['Corrected_TMAX'])
        tmin_txt  = _temp_range_text(forecast_df['Corrected_TMIN'])
        rhmax_txt = _rh_range_text(forecast_df['Corrected_RHMAX'])
        rhmin_txt = _rh_range_text(forecast_df['Corrected_RHMIN'])

        rain_txt = manual_rainfall.get(loc_name, "-a-")
        rain_txt = _NBSP if (rain_txt is None or str(rain_txt).strip().lower() == "-a-") else str(rain_txt)

        wind_txt = manual_wind.get(loc_name, "-a-")
        wind_txt = _NBSP if (wind_txt is None or str(wind_txt).strip().lower() == "-a-") else str(wind_txt)

        for r in weekly_tbl.rows:
            left = r.cells[0].text.strip()
            cls = _classify_label(left)
            if not cls:
                continue

            new_label = OUTPUT_LABELS_MARATHI.get(cls)
            if new_label:
                set_cell_text_preserve_format(r.cells[0], new_label)

            if cls == "tmax":
                set_cell_text_preserve_format(r.cells[1], tmax_txt)
            elif cls == "tmin":
                set_cell_text_preserve_format(r.cells[1], tmin_txt)
            elif cls == "rhmax":
                set_cell_text_preserve_format(r.cells[1], rhmax_txt)
            elif cls == "rhmin":
                set_cell_text_preserve_format(r.cells[1], rhmin_txt)
            elif cls == "rain":
                set_cell_text_preserve_format(r.cells[1], rain_txt)
            elif cls == "wind":
                set_cell_text_preserve_format(r.cells[1], wind_txt)

# =========================
# ===== PRUNE & SAVE ======
# =========================

def keep_only_blocks(doc: Document, keep_ranges):
    body = doc._element.body

    def _el_idx(elem):
        return elem.getparent().index(elem)

    ranges = []
    for (s, e) in keep_ranges:
        s_shift = max(0, s - 1)
        s_el = doc.paragraphs[s_shift]._element if s_shift < len(doc.paragraphs) else doc.paragraphs[-1]._element
        e_el = doc.paragraphs[e - 1]._element   if e - 1   < len(doc.paragraphs) else doc.paragraphs[-1]._element
        ranges.append((_el_idx(s_el), _el_idx(e_el)))

    def inside(elem):
        pos = _el_idx(elem)
        for a, b in ranges:
            if a <= pos <= b:
                return True
        return False

    to_remove = []
    for child in list(body):
        if child.tag.endswith('}p') or child.tag.endswith('}tbl'):
            if not inside(child):
                to_remove.append(child)
    for el in to_remove:
        body.remove(el)

    try:
        if doc.sections:
            doc.sections[-1].start_type = WD_SECTION.CONTINUOUS
    except Exception:
        pass

# =========================
# ===== FILE NAMES ========
# =========================

def single_loc_filename(loc_name, start_dt):
    return f"{loc_name}_{start_dt.day:02d}_{start_dt.month:02d}_{start_dt.year}.docx"


def group_filename(group_name, start_dt):
    return f"{group_name}_{start_dt.day:02d}_{start_dt.month:02d}_{start_dt.year}.docx"

# =========================
# ====== GENERATION  ======
# =========================

def _files_present(info: dict) -> bool:
    return os.path.exists(info["obs_file"]) and os.path.exists(info["model_file"])

def build_outputs():
    Path(OUTPUT_DIR).mkdir(parents=True, exist_ok=True)

    probe = Document(TEMPLATE_DOC_PATH)
    _, template_order = build_location_blocks(probe)

    loc_map = {x["name"]: x for x in locations}
    forecast_cache: dict[str, pd.DataFrame] = {}

    def _usable(name: str) -> bool:
        if name not in loc_map:
            return False
        info = loc_map[name]
        if not _files_present(info):
            print(f"⏭️  Skipping {name}: file(s) missing.")
            return False
        if name in forecast_cache:
            return not forecast_cache[name].empty
        try:
            fc = compute_forecast_for_location(info["obs_file"], info["model_file"], FORECAST_START, FORECAST_END)
        except Exception as e:
            print(f"⏭️  Skipping {name}: forecast error: {e}")
            return False
        if fc is None or fc.empty:
            print(f"⏭️  Skipping {name}: empty forecast.")
            return False
        forecast_cache[name] = fc
        return True

    def _finalize_and_save(doc: Document, out_path: str, keep_ranges):
        keep_only_blocks(doc, keep_ranges)
        doc.save(out_path)
        print(f"✅ Wrote: {out_path}")

    def _build_one_doc(name_list, out_path):
        doc = Document(TEMPLATE_DOC_PATH)
        blocks, _ = build_location_blocks(doc)
        keep_ranges = []

        for nm in name_list:
            if nm not in blocks:
                print(f"⚠️ Block not found in template for '{nm}', skipping…")
                continue
            fc = forecast_cache.get(nm)
            if fc is None or fc.empty:
                print(f"⏭️  Skipping '{nm}' inside document: no forecast.")
                continue

            s, e = blocks[nm]
            bulletin = int(bulletin_numbers.get(nm, 1))
            mr = manual_rainfall.get(nm, "-a-")
            mw = manual_wind.get(nm, "-a-")

            fill_location_block(doc, (s, e), nm, bulletin, fc, FORECAST_START, FORECAST_END, mr, mw)
            keep_ranges.append((s, e))

        if keep_ranges:
            _finalize_and_save(doc, out_path, keep_ranges)
        else:
            print(f"⚠️ No valid blocks for {out_path}. Nothing written.")

    if GENERATE_MODE == "per_location":
        for name in template_order:
            if _usable(name):
                out_path = Path(OUTPUT_DIR) / single_loc_filename(name.replace(" ", "_"), FORECAST_START)
                _build_one_doc([name], str(out_path))
    else:
        for out_name, name_list in GROUPS.items():
            ordered_usable = [nm for nm in template_order if nm in name_list and _usable(nm)]
            if not ordered_usable:
                print(f"⏭️  Skipping group '{out_name}': no usable locations.")
                continue
            out_path = Path(OUTPUT_DIR) / group_filename(out_name.replace(" ", "_"), FORECAST_START)
            _build_one_doc(ordered_usable, str(out_path))

# =========================
# ====== RUN IT!     ======
# =========================

if __name__ == "__main__":
    build_outputs()
